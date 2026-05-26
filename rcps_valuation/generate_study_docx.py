# -*- coding: utf-8 -*-
"""RCPS 평가 툴 스터디 자료 DOCX 생성"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ACCENT = RGBColor(0x1A, 0x56, 0xB6)   # 우리 툴 컬러
GRAY = RGBColor(0x55, 0x55, 0x55)
HEADING_BG = "EBF2FB"

def set_font(run, name="맑은 고딕", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color: run.font.color.rgb = color

def H(doc, text, level=1, color=None):
    """Heading wrapper"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level==1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    sizes = {1: 18, 2: 14, 3: 12}
    set_font(r, size=sizes.get(level, 11), bold=True, color=color or ACCENT)

def P(doc, text, bold=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p

def BULLET(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    r = p.add_run(text)
    set_font(r, size=10.5)
    return p

def TABLE_2COL(doc, header, rows, widths=(4.0, 11.0)):
    """2-column data table"""
    t = doc.add_table(rows=1+len(rows), cols=2)
    t.style = 'Light Grid Accent 1'
    # header
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    # data
    for ri, row in enumerate(rows):
        cells = t.rows[1+ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(val))
            set_font(r, size=10, bold=(ci==0))
    # width
    for col, w in zip(t.columns, widths):
        for cell in col.cells:
            cell.width = Cm(w)
    return t


doc = Document()
# 페이지 마진
sec = doc.sections[0]
sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)

# ─────────── 표지 ───────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(100)
r = title_p.add_run("RCPS 공정가치 평가 툴\n")
set_font(r, size=24, bold=True, color=ACCENT)
r = title_p.add_run("스터디 자료 — 이론·구조·활용")
set_font(r, size=14, bold=False, color=GRAY)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_before = Pt(40)
r = sub_p.add_run("K-IFRS 1109 / 1032 기준 상환전환우선주 공정가치 평가\n")
set_font(r, size=11, color=GRAY)
r = sub_p.add_run("4모형 비교 분석 (TF · GS · MC · BDT)")
set_font(r, size=11, color=GRAY)

author_p = doc.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_p.paragraph_format.space_before = Pt(180)
r = author_p.add_run("작성: 김웅(웅계사)\n")
set_font(r, size=10, color=GRAY)
from datetime import date
r = author_p.add_run(f"문서 생성일: {date.today().strftime('%Y-%m-%d')}")
set_font(r, size=10, color=GRAY)

doc.add_page_break()

# ─────────── 목차 ───────────
H(doc, "목차", level=1)
toc = [
    "1. RCPS 개요 및 회계 처리",
    "2. 평가 툴 구성 — 입력·모형·출력",
    "3. 4가지 평가 모형 비교",
    "4. 이항트리 구조와 노드별 의사결정",
    "5. 3-component 분해 (순채권·풋옵션·전환권)",
    "6. 부트스트래핑 — 무위험·신용율 곡선",
    "7. 변동성 산정 (유사기업 분석)",
    "8. 희석 효과 (전환비율·흡수형 컨벤션)",
    "9. 만기 시 처리 (풋 보장수익률·우선배당)",
    "10. 후속측정 및 민감도분석",
    "11. 감사조서 출력 및 활용",
    "12. 사용 흐름 요약",
]
for item in toc:
    P(doc, item, size=11)

doc.add_page_break()

# ─────────── 1. RCPS 개요 ───────────
H(doc, "1. RCPS 개요 및 회계 처리", level=1)

H(doc, "1.1 RCPS란?", level=2)
P(doc, "상환전환우선주(Redeemable Convertible Preferred Stock, RCPS)는 다음 세 권리를 결합한 하이브리드 증권입니다:")
BULLET(doc, "상환권(R): 발행자 또는 보유자가 약정 시점에 원금을 상환받을 권리")
BULLET(doc, "전환권(C): 보유자가 보통주로 전환할 수 있는 권리")
BULLET(doc, "우선주(P): 보통주보다 우선하여 배당을 받을 권리")
P(doc, "법적 형식은 우선주(자본)이지만, 풋옵션·강제상환 조항이 있으면 경제적 실질은 부채에 가깝습니다.", color=GRAY, size=10)

H(doc, "1.2 K-IFRS 분류 (IAS 32)", level=2)
TABLE_2COL(doc, ["분류 기준", "내용"], [
    ("부채 분류", "발행자가 현금 지급 의무(강제 상환·보유자 풋)를 가지면 → 부채"),
    ("자본 분류", "발행자에게 무조건적 의무가 없으면 → 자본"),
    ("복합금융상품", "전환권은 자본, 채권 부분은 부채로 분리 가능"),
    ("실무 표준", "한국 RCPS는 대부분 부채 분류 (풋옵션 존재)"),
])

H(doc, "1.3 K-IFRS 측정 (IFRS 9 / 1109)", level=2)
P(doc, "공정가치 측정 표준:")
BULLET(doc, "최초 인식: 공정가치 (보통 발행가액)")
BULLET(doc, "후속 측정 (FVTPL): 매 보고일 공정가치 재산정 → 평가손익 인식")
BULLET(doc, "복잡한 옵션 구조 → 이항모형 등 평가기법으로 공정가치 산정")
P(doc, "본 툴은 K-IFRS 1109 후속측정 또는 매매·취득 시 공정가치 평가를 보조합니다.",
  color=GRAY, size=10)

doc.add_page_break()

# ─────────── 2. 평가 툴 구성 ───────────
H(doc, "2. 평가 툴 구성", level=1)

H(doc, "2.1 입력 항목 — 발행 조건", level=2)
TABLE_2COL(doc, ["항목", "설명"], [
    ("발행일·만기일·평가기준일", "잔존만기(T) 산출 기초"),
    ("발행가액 (face_value)", "총 발행금액 (원). per RCPS가 아닌 총액"),
    ("전환가액 (conversion_price)", "1주당 전환가격 (원)"),
    ("RCPS 주식수 / 보통주 주식수", "희석 효과 산정용"),
    ("우선배당률 (coupon_rate)", "연 배당률 (예: 2%). 누적·비누적 선택"),
    ("풋 행사일 (put_start)", "투자자 풋 시작일"),
    ("풋 보장수익률 (put_irr)", "연 복리 IRR (예: 7.5%)"),
    ("풋 행사가 산식", "fixed / irr_y(make-whole) / cp_y(단순복리) 등"),
    ("콜 행사일·콜 IRR (선택)", "발행자 콜 옵션"),
    ("리픽싱 조건", "전환가 조정 트리거·하한·빈도"),
    ("주가·변동성·무위험율·신용스프레드", "시장 데이터"),
])

H(doc, "2.2 출력 결과", level=2)
BULLET(doc, "공정가치 (Fair Value) — 4개 모형 비교")
BULLET(doc, "3-component 분해: 순채권가치·풋옵션가치·전환권가치")
BULLET(doc, "노드별 이항트리: 주가·의사결정·전환내재가치·채권내재가치·실현가치·보유가치")
BULLET(doc, "민감도 분석 (주가·변동성·할인율)")
BULLET(doc, "후속측정 (분기말 등 시점별 재산정)")
BULLET(doc, "감사조서 Excel 자동 생성")

doc.add_page_break()

# ─────────── 3. 4가지 평가 모형 ───────────
H(doc, "3. 4가지 평가 모형 비교", level=1)

P(doc, "본 툴은 4개의 독립 평가 모형을 동시에 산출하여 결과를 교차검증합니다.")

H(doc, "3.1 Tsiveriotis-Fernandes (TF) — 주채택 모형", level=2)
P(doc, "복합금융상품 분리 평가의 학술 표준입니다.")
BULLET(doc, "지분 컴포넌트 E: 전환 시 받을 주식 가치 — 무위험율(Rf)로 할인")
BULLET(doc, "채권 컴포넌트 B: 쿠폰·상환금 — 신용조정율(Kd)로 할인")
BULLET(doc, "각 노드에서 max(전환, 풋, 보유) 의사결정 후 컴포넌트별 할인")
BULLET(doc, "K-IFRS 1109 부채/자본 분리 분류 회계처리와 정합")

H(doc, "3.2 Goldman Sachs (GS) — 블렌딩 할인", level=2)
P(doc, "5803 한국 평가실무 표준. 단일 가치(V)를 위험중립 블렌딩 할인계수로 산정:")
BULLET(doc, "Pass 1: TF식 E/B 분리로 노드별 의사결정 그리드 작성")
BULLET(doc, "Pass 2: 앞방향 전환확률 cp(t,j) 산정")
BULLET(doc, "Pass 3: 부모노드 cp로 가중평균한 단일 할인율로 후방귀납  (r_blended = cp·Rf + (1−cp)·Kd)")
P(doc, "5803/한국 회계법인 평가서들이 가장 많이 사용하는 산식.", color=GRAY, size=10)

H(doc, "3.3 Monte Carlo (MC) — 경로 시뮬레이션", level=2)
P(doc, "수만 개의 주가 경로를 시뮬레이션해 각 경로의 페이오프를 평균:")
BULLET(doc, "GBM(기하 브라운운동)으로 주가 경로 생성 (Antithetic Variates)")
BULLET(doc, "Longstaff-Schwartz Method (LSM): 회귀로 조기 행사 의사결정")
BULLET(doc, "경로의존 옵션 처리 가능: VWAP 리픽싱, 강제전환 KO, 소프트콜, 배리어 풋")
BULLET(doc, "감사 재현성: 평가일 기준 시드 고정 (동일 입력 → 동일 결과)")

H(doc, "3.4 Black-Derman-Toy (BDT) — 금리트리 교차검증", level=2)
P(doc, "독립적인 금리트리로 채권·풋채권을 평가해 교차 검증:")
BULLET(doc, "단기금리 이항트리 — 시장 zero 곡선에 캘리브레이션")
BULLET(doc, "금리 변동성(rate_vol) 반영 → 풋옵션의 시간가치 정확 산정")
BULLET(doc, "TF의 결정론적 할인 결과와 비교 → 보통 ±1% 일치하면 OK")

doc.add_page_break()

# ─────────── 4. 이항트리 구조 ───────────
H(doc, "4. 이항트리 구조와 노드별 의사결정", level=1)

H(doc, "4.1 CRR (Cox-Ross-Rubinstein) 트리", level=2)
P(doc, "주가가 매 시간격 dt마다 상승(u배) 또는 하락(d배):")
BULLET(doc, "u = exp(σ·√dt),  d = 1/u")
BULLET(doc, "위험중립확률 p = (exp((Rf−q)·dt) − d) / (u − d)")
BULLET(doc, "각 노드 (i, j): 주가 = S₀·u^(i−j)·d^j  (i=시간스텝, j=하락횟수)")
P(doc, "트리 수: 보통 월별 (T×12), 최대 120 step (5년 RCPS 기준)", color=GRAY, size=10)

H(doc, "4.2 노드별 의사결정 우선순위", level=2)
P(doc, "각 노드에서 RCPS 보유자는 다음 중 가장 유리한 선택을 함:")
TABLE_2COL(doc, ["결정", "조건"], [
    ("전환 (Convert)", "전환내재가치 > 보유가치·풋가  (i ≥ 전환 시작 step)"),
    ("풋 행사 (Redeem)", "풋가 > 보유가치·전환가치  (i ≥ 풋 시작 step)"),
    ("보유 (Hold)", "전환·풋 모두 불리 → 다음 노드까지 계속 보유"),
    ("콜 (Call, 발행자)", "발행자 콜가 < 보유가치  (i ≥ 콜 시작 step)"),
])

H(doc, "4.3 트리에 표시되는 가치 종류", level=2)
TABLE_2COL(doc, ["그리드", "의미"], [
    ("주가 트리 (stock)", "시간별 주가 (per share, 원/주)"),
    ("의사결정 (decision)", "전환/상환/콜/보유 라벨"),
    ("전환 내재가치 (conv_intrinsic)", "이 노드에서 즉시 전환 시 받을 금액"),
    ("채권 내재가치 (bond_intrinsic)", "이 노드에서 즉시 풋·만기상환 시 받을 금액"),
    ("RCPS 가치 (rcps_value)", "의사결정 후 채택된 노드 가치 (총 가치)"),
    ("보유가치 (hold_value)", "다음 노드까지 보유 시 가치"),
    ("지분가치(실현) (equity_comp)", "지분 컴포넌트 — 전환 시만 양수"),
    ("채권가치(실현) (bond_comp)", "채권 컴포넌트 — 상환·보유 시 양수"),
    ("지분·채권 보유가치", "결정과 무관 계산 — 미래 가치의 할인"),
])

doc.add_page_break()

# ─────────── 5. 3-component 분해 ───────────
H(doc, "5. 3-component 분해 — 가치 구성 요소 분석", level=1)

P(doc, "K-IFRS 1109.B4.3.5 흡수형 분해 방식 (한국 평가실무 표준):")

H(doc, "5.1 분해 산식", level=2)
TABLE_2COL(doc, ["구성요소", "산식 / 의미"], [
    ("① 순채권가치 (Bond)", "전환·풋 옵션 없는 채권 PV (만기상환 + 쿠폰을 Kd로 할인). 발행자의 무조건 의무 PV."),
    ("② 풋옵션가치 (Put)", "풋채권가치 − 순채권가치 = 조기 풋 행사로 얻는 추가 시간가치"),
    ("③ 전환권가치 (Conv)", "총 공정가치 − 풋채권가치 = 전환 선택권으로 얻는 추가 가치"),
    ("공정가치 합계", "FV = Bond + Put + Conv  (항등식 — 모든 노드 정합)"),
])

H(doc, "5.2 회계적 의의", level=2)
BULLET(doc, "① 순채권: 부채 부분 (이자비용 인식 기초)")
BULLET(doc, "② 풋옵션: 부채 부분에 흡수 (행사 강제성 있으면)")
BULLET(doc, "③ 전환권: 자본 부분 (자본 분류 시) / 또는 파생상품 부채")
P(doc, "K-IFRS 1109.B4.3.5: 복합금융상품의 부채 부분 = 발행자의 무조건적 의무.\n"
      "전환권은 부채 분리 후 잔여 가치로 자본 또는 파생상품 부채로 분류.",
  color=GRAY, size=10)

H(doc, "5.3 5803 케이스 예시", level=2)
TABLE_2COL(doc, ["항목", "값"], [
    ("총 공정가치", "120,467,994,612원"),
    ("① 순채권가치", "78,003,886,158원 (64.8%)"),
    ("② 풋옵션가치", "6,423,378,864원 (5.3%)"),
    ("③ 전환권가치", "36,040,729,590원 (29.9%)"),
])

doc.add_page_break()

# ─────────── 6. 부트스트래핑 ───────────
H(doc, "6. 부트스트래핑 — 무위험율·신용율 기간구조", level=1)

H(doc, "6.1 왜 곡선이 필요한가?", level=2)
P(doc, "평탄(flat) 단일 금리는 만기 시점별 다른 금리 환경을 반영 못 함:")
BULLET(doc, "단기금리와 장기금리가 다른 시장 상황 (yield curve)")
BULLET(doc, "신용율도 만기별 다름 (회사채 신용곡선)")
BULLET(doc, "노드별 forward rate로 할인하면 시장가격에 더 정합")

H(doc, "6.2 입력 데이터 — 두 가지 곡선", level=2)
TABLE_2COL(doc, ["곡선", "출처"], [
    ("RF (무위험)", "국고채 YTM (ECOS·KIS-NET) → 부트스트랩으로 스팟·forward 산출"),
    ("RD (신용조정)", "회사채 신용등급별 YTM (KIS-NET) → 동일 부트스트랩"),
    ("신용스프레드", "RD − RF (만기별 차이)"),
])

H(doc, "6.3 부트스트래핑 산식", level=2)
P(doc, "YTM(연복리) → D-SPOT(연이산스팟) → C-SPOT(연속스팟) → C-FWD(연속선도) 단계 변환:")
BULLET(doc, "D-SPOT z_d(t): zero coupon 채권의 현재가가 시장가와 일치하는 스팟율")
BULLET(doc, "C-SPOT z_c(t) = ln(1 + z_d(t)) — 연속복리 변환")
BULLET(doc, "C-FWD f(t1, t2) = (z_c(t2)·t2 − z_c(t1)·t1) / (t2 − t1)")
BULLET(doc, "각 step별 forward rate를 이항트리 할인계수로 사용")

doc.add_page_break()

# ─────────── 7. 변동성 ───────────
H(doc, "7. 변동성 산정", level=1)

H(doc, "7.1 변동성의 역할", level=2)
P(doc, "변동성 σ는 옵션 가치의 핵심 요소:")
BULLET(doc, "주가 트리의 u/d 결정 → 미래 주가 분포 폭")
BULLET(doc, "옵션 가치는 σ의 증가함수 (변동성 ↑ → 전환권 가치 ↑)")
BULLET(doc, "감사 시 변동성 검증이 가장 민감 — 변동성 1%p 차이가 FV 수~수십% 차이 야기")

H(doc, "7.2 유사기업 비교 분석", level=2)
P(doc, "비상장 RCPS 평가 시 유사 상장기업의 역사적 변동성을 활용:")
BULLET(doc, "Yahoo Finance / KRX에서 일별 주가 시계열 (보통 1~3년)")
BULLET(doc, "일일 로그수익률 표준편차 × √252 (영업일 기준) → 연간 변동성")
BULLET(doc, "유사기업 다수 선정 → 중앙값 또는 평균 사용")
BULLET(doc, "이상치(outlier) 필터링 — IQR 또는 MAD 방법")

H(doc, "7.3 본 툴의 변동성 산출 흐름", level=2)
BULLET(doc, "사용자가 평가대상 + 유사기업 티커 입력")
BULLET(doc, "각 기업의 일별 종가 자동 다운로드 (KRX / Yahoo / NASDAQ 등)")
BULLET(doc, "로그수익률 시계열 → 표본 표준편차 → 연환산 변동성")
BULLET(doc, "유사기업 변동성의 평균/중앙값 → 평가 대상에 적용")

doc.add_page_break()

# ─────────── 8. 희석 ───────────
H(doc, "8. 희석 효과 — 전환비율 및 흡수형 컨벤션", level=1)

H(doc, "8.1 전환 시 신주 발행 영향", level=2)
P(doc, "RCPS가 보통주로 전환되면 새로운 주식이 발행되어 기존 보통주 가치가 희석됩니다:")
BULLET(doc, "전환 비율 = face_per_RCPS / conversion_price (보통 1:1)")
BULLET(doc, "전환 후 총 주식수 = 기존 보통주 + 전환된 신주")
BULLET(doc, "per-share 가치 ↓ (분모 증가)")

H(doc, "8.2 두 가지 평가 컨벤션", level=2)
TABLE_2COL(doc, ["컨벤션", "산식 / 적용 케이스"], [
    ("비-희석 (단순)", "전환가치 = (face/K) × stock_price.\n주가가 시장에서 이미 dilution 가능성을 반영한다고 가정."),
    ("희석 (흡수형)", "전환 시 회사가치 = 기존지분(n_com·S) + 흡수된 부채(bond_pv).\n전환 후 1주당 가치 = (n_com·S + bond_pv) / (n_com + N_new).\n전환가치 = 1주당 × N_new (=face/K, 전환 시 발행 신주 총수)."),
])

H(doc, "8.3 어느 쪽을 선택할까?", level=2)
BULLET(doc, "발행 규모가 시총 대비 작으면 → 비-희석으로 충분 (실무 관행)")
BULLET(doc, "발행 규모가 크고 희석률 ≥ 5% → 희석 모형 적용 권장")
BULLET(doc, "한국 5803 등 표준 평가서 → 비-희석 컨벤션이 다수")
BULLET(doc, "본 툴은 두 모드 모두 지원 — common_shares 입력 여부로 자동 분기")

doc.add_page_break()

# ─────────── 9. 만기 처리 ───────────
H(doc, "9. 만기 시 처리 — 풋 보장수익률·우선배당", level=1)

H(doc, "9.1 풋 보장수익률 (IRR)", level=2)
P(doc, "투자자에게 약정된 최소 수익률. 만기 풋 행사 시 받는 금액:")
BULLET(doc, "고정가 (fixed): 약정된 정액")
BULLET(doc, "cp_y (단순 복리): face × (1 + IRR)^t")
BULLET(doc, "irr_y (make-whole): face × (1 + IRR)^t − 누적 우선배당의 FV (이중 보상 방지)")
P(doc, "irr_y는 한국 RCPS 표준 — 풋 IRR과 우선배당이 이중 보상되지 않도록 차감.",
  color=GRAY, size=10)

H(doc, "9.2 우선배당 — 만기까지의 처리", level=2)
P(doc, "RCPS 우선배당은 다음 방식으로 처리:")
BULLET(doc, "누적성 (cumulative): 미지급분 누적 → 다음 지급일에 일괄")
BULLET(doc, "첫 지급 시점 (dividend_first_pay_year): 발행 후 N년부터 지급 가능 (배당가능이익 가정)")
BULLET(doc, "전환 시 처리: 보통 누적 미지급분은 지급 후 보통주 전환 (한국 표준)")

H(doc, "9.3 만기 시 의사결정 — TF vs GS 차이", level=2)
TABLE_2COL(doc, ["모형", "만기 컨벤션"], [
    ("TF (이론적 표준)", "전환 시 만기쿠폰 포기 (forfeit) — 보수적"),
    ("GS / 한국 실무", "전환·상환 무관 만기쿠폰 지급 — 흡수형 컨벤션"),
    ("권고", "약정서에 '전환 시 누적 미지급 우선배당 일시 지급' 조항 있으면 GS 사용"),
])

doc.add_page_break()

# ─────────── 10. 후속측정·민감도 ───────────
H(doc, "10. 후속측정 및 민감도 분석", level=1)

H(doc, "10.1 후속측정 (Subsequent Measurement)", level=2)
P(doc, "K-IFRS 1109 FVTPL 후속측정: 매 보고일 공정가치 재산정.")
BULLET(doc, "각 보고일 시점의 시장 데이터 입력 (주가·변동성·금리 곡선)")
BULLET(doc, "최초 인식과 동일한 평가 모형·할인 컨벤션 유지 (측정 연속성)")
BULLET(doc, "전기 대비 변동: 평가손익 (당기손익 또는 OCI)")
BULLET(doc, "본 툴: 분기말·반기말·연도말 등 다중 시점 일괄 산정 가능")

H(doc, "10.2 민감도 분석", level=2)
P(doc, "주요 입력 변수의 변화에 대한 FV 민감도 측정:")
TABLE_2COL(doc, ["변수", "민감도 범위"], [
    ("주가 (S)", "±20% 범위 — 5% 단위"),
    ("변동성 (σ)", "±50% 범위 — 10% 단위"),
    ("무위험율 (Rf)", "±100bp — 25bp 단위"),
    ("신용스프레드", "±100bp — 25bp 단위"),
])
P(doc, "감사 시 결과의 견고성(robustness) 확인용. 핵심 가정 범위 내 FV 변동성 검토.",
  color=GRAY, size=10)

doc.add_page_break()

# ─────────── 11. 감사조서 ───────────
H(doc, "11. 감사조서 출력", level=1)

P(doc, "본 툴은 K-IFRS 1109 평가 감사조서를 Excel 파일로 자동 생성합니다.")

H(doc, "11.1 감사조서 시트 구성", level=2)
TABLE_2COL(doc, ["시트", "내용"], [
    ("1. 발행조건", "약정 조건 전체 (일자·금액·이자율·옵션 조건)"),
    ("2. 평가결과", "TF 주채택 FV + 이항 파라미터 (u/d/p/steps)"),
    ("3. 모형비교", "TF·GS·MC 모형별 결과 (Bond·Put·Conv·FV)"),
    ("4. BDT교차검증", "TF vs BDT 채권 비교 — Δ% 표시"),
    ("5. TF이항트리", "9개 그리드 (주가·의사결정·내재가치·실현·보유)"),
    ("6. GS이항트리", "6개 그리드 (주가·의사결정·전환확률·할인계수·내재가치·가치)"),
    ("7. 후속측정", "분기말 등 보고일별 FV·평가손익"),
    ("8. 민감도분석", "주가·변동성·금리 민감도"),
])

H(doc, "11.2 감사 활용", level=2)
BULLET(doc, "외부 감사인의 RCPS 공정가치 평가 적정성 검증 자료")
BULLET(doc, "K-IFRS 1109 부합성 입증 — 표준 모형·표준 컨벤션 사용")
BULLET(doc, "이항트리로 노드별 의사결정 추적 가능 (의사결정 색상 코딩)")
BULLET(doc, "다중 모형 결과 제시로 모형 위험 완화")
BULLET(doc, "감사 재현성: 동일 입력 시 동일 결과 (시드 고정 - K-IFRS 13.91)")

doc.add_page_break()

# ─────────── 12. 사용 흐름 ───────────
H(doc, "12. 사용 흐름 요약", level=1)

H(doc, "STEP 1 — 시장 데이터 준비", level=2)
BULLET(doc, "부트스트래핑 탭: ECOS 국고채 + KIS-NET 회사채 → Rf·Rd 곡선 생성")
BULLET(doc, "변동성 탭: 평가대상 + 유사기업 티커 입력 → 연환산 변동성 자동 산정")

H(doc, "STEP 2 — 발행조건 입력", level=2)
BULLET(doc, "RCPS 공정가치 평가 탭: 발행일·만기·금액·전환가·풋·콜 조건 입력")
BULLET(doc, "리픽싱 조건 (있는 경우): 트리거·하한·빈도")
BULLET(doc, "강제전환·소프트콜·배리어풋 등 경로의존 옵션 (있는 경우)")

H(doc, "STEP 3 — 평가 실행", level=2)
BULLET(doc, "'공정가치 산출' 버튼 → 4개 모형 동시 평가 (TF·GS·MC·BDT)")
BULLET(doc, "결과 카드: 3-component 분해·모형별 FV·차이 분석")
BULLET(doc, "이항트리 시각화: 노드별 의사결정·가치 (TF/GS 선택 가능)")

H(doc, "STEP 4 — 검증 및 보고", level=2)
BULLET(doc, "BDT 교차검증: 단기금리 변동성 입력 → 채권 PV 독립 검증")
BULLET(doc, "민감도 분석: 핵심 가정 변동 범위 내 FV 변동성")
BULLET(doc, "후속측정 (필요시): 분기말 등 시점별 재산정")
BULLET(doc, "감사조서 Excel 다운로드: K-IFRS 1109 부합 문서화")

# ─────────── 마무리 ───────────
H(doc, "참고 자료", level=1)
BULLET(doc, "Tsiveriotis & Fernandes (1998): Valuing Convertible Bonds with Credit Risk")
BULLET(doc, "Goldman Sachs (1994): Valuing Convertible Bonds as Derivatives")
BULLET(doc, "Black, Derman, Toy (1990): One-Factor Model of Interest Rates")
BULLET(doc, "Longstaff, Schwartz (2001): Valuing American Options by Simulation")
BULLET(doc, "K-IFRS 1109 (IFRS 9) 금융상품 — 부채/자본 분류, 측정")
BULLET(doc, "K-IFRS 1032 (IAS 32) 금융상품: 표시")
BULLET(doc, "K-IFRS 1113 (IFRS 13) 공정가치 측정")

# 저장
out_path = os.path.expanduser('~/Downloads/RCPS_평가툴_스터디자료.docx')
doc.save(out_path)
print(f'생성 완료: {out_path}')
print(f'크기: {os.path.getsize(out_path):,} bytes')
